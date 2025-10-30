from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager
import time
from docx import Document
import requests
from docx.shared import Inches


# Setup
# driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()))
chrome_options = webdriver.ChromeOptions()

# If browser version error: enable the line below and change "141" to the version you have currently now
# chrome_options.browser_version = "141"

driver = webdriver.Chrome(options=chrome_options)
driver.set_window_size(1024, 668)
driver.get("https://web.slowly.app")

#  Wait for user to log in manually
input("Please log in to Slowly Web and press Enter here...")

# Navigate to conversation
friend_name = driver.find_element(By.CSS_SELECTOR, "#root > div > div.main-scroller > div > div > div.pl-3.main-container.flex-grow-1 > div > div.friend-header-wrapper > div > div > div.col-9.pt-2 > span").text

# Scroll to load all letters
SCROLL_PAUSE_TIME = 2
last_height = driver.execute_script("return document.body.scrollHeight")

while True:
    driver.execute_script("window.scrollTo(0, document.body.scrollHeight);")
    time.sleep(SCROLL_PAUSE_TIME)
    new_height = driver.execute_script("return document.body.scrollHeight")
    if new_height == last_height:
        break
    last_height = new_height

# Locate letters in the conversation
letter_elements = driver.find_elements(By.CSS_SELECTOR, "#root .main-scroller .main-container.flex-grow-1 .friend-letters-wrapper.p-2 > div > div.col-6.col-xl-4.mb-3")
# letter_elements = driver.find_elements(By.CLASS_NAME, "col-6 col-xl-4 mb-3")
print('amount of letters:', len(letter_elements))

# Create Word document
doc = Document('slowly_letters_template.docx')
doc.add_heading(f"Letters with {friend_name}", 0)

for i in range(len(letter_elements)):
    try:
        cards = driver.find_elements(By.CSS_SELECTOR, "#root .main-scroller .main-container.flex-grow-1 .friend-letters-wrapper.p-2 > div > div.col-6.col-xl-4.mb-3")
        print(f"Opening letter {i+1}...")
        card = cards[i]
        # print(card)
        card.click()
        time.sleep(2)  # Allow letter content to load

        # date:
        date = driver.find_element(By.CSS_SELECTOR, "#root > div > div.main-scroller > div > div > div.pl-3.main-container.flex-grow-1 > div > div.sticky-top.friend-header-wrapper > div > div > div.col.ml-n1 > div:nth-child(2) > span:nth-child(1)").text
        author = driver.find_element(By.CSS_SELECTOR, "#root > div > div.main-scroller > div > div > div.pl-3.main-container.flex-grow-1 > div > div.sticky-top.friend-header-wrapper > div > div > div.col.ml-n1 > div.text-primary.h6").text
        # Now extract content inside opened letter
        content = driver.find_element(By.CSS_SELECTOR, "#root > div > div.main-scroller > div > div > div.pl-3.main-container.flex-grow-1 > div > div.friend-Letter-wrapper.p-3.pl-4.pr-4 > div > div.modal-body > div").text
        doc.add_heading(f"{author} - {date}", level=2)
        doc.add_paragraph(content)    

        # Check for an image attachments  
        try:
            image = driver.find_element(By.CSS_SELECTOR, "#root > div > div.main-scroller > div > div > div.pl-3.main-container.flex-grow-1 > div > div.friend-Letter-wrapper.p-3.pl-4.pr-4 > div > div.modal-body > div.slider > div > img")
            img_url = image.get_attribute("src")
            img_data = requests.get(img_url).content
            img_filename = f"image_{author}_{i}.jpg"
            with open(img_filename, "wb") as f:
                f.write(img_data)
            doc.add_picture(img_filename,  width = Inches(4))  
        except Exception as e1:
            pass
        # If there are more than one image attached..
        try:
            image = driver.find_element(By.CSS_SELECTOR, "#root > div > div.main-scroller > div > div > div.pl-3.main-container.flex-grow-1 > div > div.friend-Letter-wrapper.p-3.pl-4.pr-4 > div > div.modal-body > div.slider > div > div > div > div.slick-slide.slick-active.slick-current > div > div > img")
            image_list = []
            next_image_button = driver.find_element(By.CSS_SELECTOR, "#root > div > div.main-scroller > div > div > div.pl-3.main-container.flex-grow-1 > div > div.friend-Letter-wrapper.p-3.pl-4.pr-4 > div > div.modal-body > div.slider > div > button.slick-arrow.slick-next")
            img_url = image.get_attribute("src")
            image_list.append(img_url)
            flag = 0
            amount = 1
            while flag == 0:
                next_image_button.click() 
                image = driver.find_element(By.CSS_SELECTOR, "#root > div > div.main-scroller > div > div > div.pl-3.main-container.flex-grow-1 > div > div.friend-Letter-wrapper.p-3.pl-4.pr-4 > div > div.modal-body > div.slider > div > div > div > div.slick-slide.slick-active.slick-current > div > div > img")
                img_url = image.get_attribute("src")
                try:
                    for j in range(amount):                      
                        if img_url == image_list[j]:
                            flag = 1
                    if flag != 1:
                        amount += 1
                        image_list.append(img_url)
                        next_image_button = driver.find_element(By.CSS_SELECTOR, "#root > div > div.main-scroller > div > div > div.pl-3.main-container.flex-grow-1 > div > div.friend-Letter-wrapper.p-3.pl-4.pr-4 > div > div.modal-body > div.slider > div > button.slick-arrow.slick-next")
                except Exception as e:
                    print(e)
                    break
            if len(image_list) > 0:
                for i in range(len(image_list)):
                    img_url = image_list[i]
                    img_data = requests.get(img_url).content
                    img_filename = f"image_{author}_{i}.jpg"
                    with open(img_filename, "wb") as f:
                        f.write(img_data)
                    doc.add_picture(img_filename, width = Inches(4))    
        except Exception as e:
            pass
        # button for the exit from letter:
        button = driver.find_element(By.CSS_SELECTOR, "#root > div > div.main-scroller > div > div > div.pl-3.main-container.flex-grow-1 > div > div.sticky-top.friend-header-wrapper > div > div > a > i")
        button.click()
        time.sleep(3)
    except Exception as e:
        print(f"Error with letter {i+1}: {e}")
        break

doc.save(f"slowly_letters_with_{friend_name}.docx")
driver.quit()